from docx import Document
import re
from dataclasses import dataclass
from typing import List
@dataclass
class ExtractionRule:
    name:str
    regex:List[str]



class RulerBasedParser:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.extraction_rules:List[ExtractionRule] = [
            ExtractionRule( name='Counterparty', regex=[r'(?i)^party\s+a$']),
            ExtractionRule(name='Initial Valuation Date', regex=[r'(?i)^initial\s+valuation\s+date$']),
            ExtractionRule(name='Notional', regex=[r'(?i)^notional\s+amount\s+\(n\)$',r'(?i)^notional\s+amount$']),
            ExtractionRule(name='Valuation Date', regex=[r'(?i)^valuation\s+date$']),
            ExtractionRule(name='Maturity', regex=[r'(?i)^termination\s+date$']),
            ExtractionRule(name='Underlying', regex=[r'(?i)^underlying$']),
            ExtractionRule(name='Coupon', regex=[r'(?i)^coupon\s+\(c\)$',r'(?i)^coupon$']),
            ExtractionRule(name='Barrier', regex=[r'(?i)^barrier\s+\(b\)$',r'(?i)^barrier$']),
            ExtractionRule(name='calendar', regex=[r'(?i)^business\s+day$']),

        ]

    def parse(self):
        for table in self.doc.tables:
            print("Rule Based Parser")
            for row in table.rows:
                # print(row.cells)
                for extraction_rule in self.extraction_rules:
                    for regex in extraction_rule.regex:
                        if re.match(regex, row.cells[0].text):
                            # print(extraction_rule.name)
                            # print(row.cells[0].text)
                            for i in range(1, len(row.cells)):
                                print(f"{extraction_rule.name} -> {row.cells[0].text} : {row.cells[i].text}")
                # else:
                    # print(f" {row.cells[0].text} is not in defined data")
                # for cell in row.cells:
                #     if cell.text in defined_data:
                #         print(cell.text)
                #     else:
                #         print(f" {cell.text} is not in defined data")
                # for cell in row.cells:
                #     print(cell.text)
        # for paragraph in self.doc.paragraphs:
        #     print(paragraph.text)

if __name__ == "__main__":
    ruler_based_parser = RulerBasedParser("ZF4894_ALV_07Aug2026_physical.docx")
    ruler_based_parser.parse()